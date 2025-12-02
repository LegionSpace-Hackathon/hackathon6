# -*- coding: utf-8 -*-
"""
Linux File Operations MCP Server
基于FastMCP框架的Linux文件操作服务器，支持远程文件管理、Python依赖安装等功能
"""

import os
import socket
import logging
import json
import time
import shutil
from typing import Optional, Dict, Any
import paramiko
import base64
from mcp.server.fastmcp import FastMCP
from fastapi import FastAPI, File, UploadFile, HTTPException, Form, Query
from fastapi.responses import JSONResponse, FileResponse, StreamingResponse
import tempfile
import uvicorn
import threading
import asyncio
import re
from pathlib import Path

# Markdown转换相关导入（可选，如果导入失败则使用备选方案）
PYPANDOC_AVAILABLE = False
WEASYPRINT_AVAILABLE = False
PYTHON_DOCX_AVAILABLE = False
MARKDOWN_AVAILABLE = False

try:
    import pypandoc
    PYPANDOC_AVAILABLE = True
except ImportError:
    pass

try:
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
except (ImportError, OSError, Exception) as e:
    WEASYPRINT_AVAILABLE = False
    # WeasyPrint在Windows上需要GTK+库，如果导入失败不影响服务启动
    pass

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    pass

try:
    import markdown
    from markdown.extensions import codehilite, tables, fenced_code
    MARKDOWN_AVAILABLE = True
except ImportError:
    pass

# 配置日志
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler('mcp_server.log', encoding='utf-8')
    ]
)
logger = logging.getLogger(__name__)

# 服务器配置
REMOTE_HOST = "192.168.94.62"  # 远程服务器IP，可以是localhost/127.0.0.1（本地与远程为同一台机器）
USERNAME = "root"
PASSWORD = "tfd@2025"
# HTTP服务器地址配置（用于生成下载URL等）
HTTP_SERVER_HOST = "192.168.94.62"  # HTTP服务器地址，用于生成下载链接
# Conda环境配置（可选，仅用于远程服务器）
# 说明：
# 1. 本地服务器（运行server.py的机器）：如果使用conda环境运行server.py，只需在conda环境中安装依赖包即可，无需配置此项
# 2. 远程服务器（REMOTE_HOST）：如果远程服务器上的Python代码部署在conda环境中，设置此环境名称
#    系统会在执行远程Python相关命令（install_python_deps、run_python_file等）时自动激活指定的conda环境
# 3. 本地与远程为同一台机器的情况：
#    - REMOTE_HOST 可以设置为 "127.0.0.1"、"localhost" 或本机实际IP
#    - 仍然需要通过SSH连接（即使是自己连自己），请确保SSH服务正常运行
#    - CONDA_ENV_NAME 仍用于通过SSH执行的远程命令，而非本地运行server.py的环境
# 如果远程服务器未使用conda环境或环境未设置，保持为空字符串即可
CONDA_ENV_NAME = "py312"  # 例如: "myenv" 或 "base"（仅用于远程服务器的conda环境）

# 监控配置：周期性输出线程状态
THREAD_STATUS_INTERVAL = 300  # 秒

# 创建MCP服务器
mcp = FastMCP("linux-file-ops-server", port=8001, host="0.0.0.0")

# 创建FastAPI应用用于HTTP接口
app = FastAPI(title="Linux File Operations API", version="1.0.0")

# 允许的文件格式
ALLOWED_FILE_EXTENSIONS = {'.xlsx', '.csv', '.xls'}
RISK_ANALYSIS_DIR = "/usr/local/risk_analysis/"

# 文档转换输出目录（本地路径）
DOCUMENT_OUTPUT_DIR = "/usr/local/risk_analysis/docs/"

def _get_conda_init_code() -> str:
    """
    获取conda环境初始化代码片段（用于在脚本中激活conda环境）
    
    返回：
        str: conda初始化代码，如果未配置conda环境则返回空字符串
    """
    if not CONDA_ENV_NAME or CONDA_ENV_NAME.strip() == "":
        return ""
    
    env_name = CONDA_ENV_NAME.strip()
    # 初始化conda并激活环境
    # 使用bash source方式初始化conda，然后激活指定环境
    conda_init = (
        "if command -v conda >/dev/null 2>&1; then "
        f"source $(conda info --base 2>/dev/null)/etc/profile.d/conda.sh 2>/dev/null || true; "
        f"conda activate {env_name} 2>/dev/null || true; "
        "fi; "
    )
    return conda_init

def _get_conda_deactivate_code() -> str:
    """
    获取conda环境退出代码片段（用于在脚本结束时退出conda环境）
    
    返回：
        str: conda退出代码，如果未配置conda环境则返回空字符串
    """
    if not CONDA_ENV_NAME or CONDA_ENV_NAME.strip() == "":
        return ""
    
    # 退出conda环境
    conda_deinit = "conda deactivate 2>/dev/null || true; "
    return conda_deinit

def _get_python_command() -> str:
    """
    获取正确的Python命令：在conda环境中使用python，否则使用python3
    
    返回：
        str: Python命令（"python" 或 "python3"）
    """
    if CONDA_ENV_NAME and CONDA_ENV_NAME.strip() != "":
        # 如果配置了conda环境，使用python命令（conda环境激活后python指向环境中的解释器）
        return "python"
    else:
        # 否则使用python3
        return "python3"

def _validate_file_format(filename: str) -> tuple[bool, str]:
    """
    验证文件格式是否为允许的表格格式
    
    参数：
        - **filename**: 文件名
    
    返回：
        tuple[bool, str]: (是否有效, 错误信息)
    """
    if not filename:
        return False, "文件名不能为空"
    
    # 获取文件扩展名
    file_ext = os.path.splitext(filename.lower())[1]
    
    if file_ext not in ALLOWED_FILE_EXTENSIONS:
        allowed_formats = ", ".join(ALLOWED_FILE_EXTENSIONS)
        return False, f"文件格式错误，只支持以下格式: {allowed_formats}"
    
    return True, ""

def _execute_remote_command(command: str) -> str:
    """执行远程命令，增加重试机制和更好的错误处理"""
    logger.info(f"=== 执行远程命令 ===")
    logger.info(f"命令: {command}")
    logger.info(f"目标主机: {REMOTE_HOST}")
    logger.info(f"用户名: {USERNAME}")
    
    max_retries = 3
    retry_delay = 1
    
    for attempt in range(max_retries):
        client = paramiko.SSHClient()
        client.set_missing_host_key_policy(paramiko.AutoAddPolicy())
        try:
            logger.info(f"尝试连接远程主机 {REMOTE_HOST} (第 {attempt + 1} 次)")
            start_time = time.time()
            client.connect(REMOTE_HOST, username=USERNAME, password=PASSWORD, timeout=30)
            connect_time = time.time() - start_time
            logger.info(f"SSH连接成功，耗时: {connect_time:.3f}秒")
            
            logger.info(f"执行命令: {command}")
            exec_start_time = time.time()
            stdin, stdout, stderr = client.exec_command(command)
            exec_time = time.time() - exec_start_time
            logger.info(f"命令执行完成，耗时: {exec_time:.3f}秒")
            
            out = stdout.read().decode("utf-8", errors="ignore")
            err = stderr.read().decode("utf-8", errors="ignore")
            
            logger.info(f"标准输出: {out}")
            if err:
                logger.warning(f"标准错误: {err}")
                return f"ERROR: {err.strip()}"
            
            logger.info(f"命令执行成功，输出长度: {len(out)}")
            return out.strip()
            
        except paramiko.AuthenticationException as e:
            logger.error(f"SSH 认证失败: {str(e)}")
            return "ERROR: SSH 认证失败，请检查用户名和密码"
        except paramiko.SSHException as e:
            logger.error(f"SSH 连接异常: {str(e)}")
            if attempt < max_retries - 1:
                logger.warning(f"SSH 连接异常，{retry_delay}秒后重试: {str(e)}")
                time.sleep(retry_delay)
                retry_delay *= 2
                continue
            return f"ERROR: SSH 连接异常: {str(e)}"
        except socket.timeout as e:
            logger.error(f"连接超时: {str(e)}")
            if attempt < max_retries - 1:
                logger.warning(f"连接超时，{retry_delay}秒后重试")
                time.sleep(retry_delay)
                retry_delay *= 2
                continue
            return "ERROR: 连接超时，请检查网络或服务器状态"
        except Exception as e:
            logger.error(f"连接失败: {str(e)}", exc_info=True)
            if attempt < max_retries - 1:
                logger.warning(f"连接失败，{retry_delay}秒后重试: {str(e)}")
                time.sleep(retry_delay)
                retry_delay *= 2
                continue
            return f"ERROR: 连接失败: {str(e)}"
        finally:
            try:
                client.close()
                logger.debug("SSH连接已关闭")
            except Exception as e:
                logger.warning(f"关闭SSH连接时出错: {str(e)}")
    
    logger.error("所有重试尝试都失败了")
    return "ERROR: 所有重试尝试都失败了"


def _sftp_write_file(
    file_path: str,
    content: str,
    mode: Optional[int] = None,
    append: bool = False,
) -> str:
    client = paramiko.SSHClient()
    client.set_missing_host_key_policy(paramiko.AutoAddPolicy())
    try:
        client.connect(REMOTE_HOST, username=USERNAME, password=PASSWORD, timeout=30)
        sftp = client.open_sftp()
        directory = os.path.dirname(file_path)
        if directory:
            # Ensure directory exists
            try:
                sftp.stat(directory)
            except FileNotFoundError:
                _execute_remote_command(f"mkdir -p {directory}")
        with sftp.file(file_path, mode=("a" if append else "w")) as f:
            f.write(content)
        if mode is not None:
            sftp.chmod(file_path, mode)
        sftp.close()
        return "OK"
    except paramiko.AuthenticationException:
        return "ERROR: SSH 认证失败，请检查用户名和密码"
    except paramiko.SSHException as e:
        return f"ERROR: SSH 连接异常: {str(e)}"
    except socket.timeout:
        return "ERROR: 连接超时，请检查网络或服务器状态"
    except Exception as e:
        return f"ERROR: 文件操作失败: {str(e)}"
    finally:
        client.close()


def _sftp_upload_file(
    local_path: str,
    remote_dir: str,
    filename: Optional[str] = None,
    mode: Optional[int] = None,
) -> str:
    client = paramiko.SSHClient()
    client.set_missing_host_key_policy(paramiko.AutoAddPolicy())
    try:
        if not os.path.isfile(local_path):
            return f"ERROR: 本地文件不存在: {local_path}"
        client.connect(REMOTE_HOST, username=USERNAME, password=PASSWORD, timeout=30)
        sftp = client.open_sftp()
        # 确保远程目录存在
        try:
            if remote_dir:
                sftp.stat(remote_dir)
        except FileNotFoundError:
            _execute_remote_command(f"mkdir -p {remote_dir}")
        # 组合远程文件完整路径
        remote_filename = filename if filename else os.path.basename(local_path)
        remote_path = os.path.join(remote_dir, remote_filename)
        # 直接使用 put 以支持二进制与大文件
        sftp.put(local_path, remote_path)
        if mode is not None:
            sftp.chmod(remote_path, mode)
        sftp.close()
        return remote_path
    except paramiko.AuthenticationException:
        return "ERROR: SSH 认证失败，请检查用户名和密码"
    except paramiko.SSHException as e:
        return f"ERROR: SSH 连接异常: {str(e)}"
    except socket.timeout:
        return "ERROR: 连接超时，请检查网络或服务器状态"
    except Exception as e:
        return f"ERROR: 文件上传失败: {str(e)}"
    finally:
        client.close()


def _sftp_upload_bytes(
    content: bytes,
    remote_dir: str,
    filename: str,
    mode: Optional[int] = None,
) -> str:
    client = paramiko.SSHClient()
    client.set_missing_host_key_policy(paramiko.AutoAddPolicy())
    try:
        if not filename:
            return "ERROR: 必须提供 filename 用于流式上传"
        client.connect(REMOTE_HOST, username=USERNAME, password=PASSWORD, timeout=30)
        sftp = client.open_sftp()
        # 确保远程目录存在
        try:
            if remote_dir:
                sftp.stat(remote_dir)
        except FileNotFoundError:
            _execute_remote_command(f"mkdir -p {remote_dir}")
        remote_path = os.path.join(remote_dir, filename)
        # 以二进制写入，避免临时文件，支持大文件分块
        with sftp.file(remote_path, mode="wb") as remote_f:
            # 分块写入，降低内存峰值
            chunk_size = 1024 * 1024
            offset = 0
            total = len(content)
            while offset < total:
                end = min(offset + chunk_size, total)
                remote_f.write(content[offset:end])
                offset = end
        if mode is not None:
            sftp.chmod(remote_path, mode)
        sftp.close()
        return remote_path
    except paramiko.AuthenticationException:
        return "ERROR: SSH 认证失败，请检查用户名和密码"
    except paramiko.SSHException as e:
        return f"ERROR: SSH 连接异常: {str(e)}"
    except socket.timeout:
        return "ERROR: 连接超时，请检查网络或服务器状态"
    except Exception as e:
        return f"ERROR: 文件上传失败: {str(e)}"
    finally:
        client.close()


# MCP工具定义
@mcp.tool()
async def create_file(file_path: str, content: str = "") -> str:
    """
    在远程 Linux 服务器创建文件，若包含内容则写入。
    
    参数：
        - **file_path**: 绝对或相对路径
        - **content**: 要写入的文本内容（可为空）
    
    返回：
        str: 操作结果字符串
    """
    logger.info(f"创建文件: {file_path}")
    # 在单独线程中执行阻塞的 SFTP 操作，避免阻塞事件循环
    result = await asyncio.to_thread(_sftp_write_file, file_path, content)
    logger.info(f"创建文件结果: {result}")
    return result


@mcp.tool()
async def write_file(file_path: str, content: str, append: bool = False) -> str:
    """
    向远程文件写入内容。
    
    参数：
        - **file_path**: 文件路径
        - **content**: 要写入的内容
        - **append**: 是否追加；False 则覆盖
    
    返回：
        str: 操作结果字符串
    """
    logger.info(f"写入文件: {file_path}, 追加模式: {append}")
    # 在单独线程中执行阻塞的 SFTP 操作
    result = await asyncio.to_thread(_sftp_write_file, file_path, content, None, append)
    logger.info(f"写入文件结果: {result}")
    return result


@mcp.tool()
async def chmod(file_path: str, permissions: str) -> str:
    """
    修改远程文件权限，例如 permissions="755"。
    
    参数：
        - **file_path**: 文件路径
        - **permissions**: 权限字符串，如 "755" 或 "644"
    
    返回：
        str: 操作结果字符串
    """
    logger.info(f"修改权限: {file_path} -> {permissions}")
    # 远程命令执行是阻塞的，放到线程池中
    result = await asyncio.to_thread(_execute_remote_command, f"chmod {permissions} {file_path}")
    logger.info(f"修改权限结果: {result}")
    return result


@mcp.tool()
async def install_python_deps(target_path: str, index_url: Optional[str] = None, force_regenerate: bool = False) -> str:
    """
    在远程服务器为指定的 Python 文件或目录安装依赖。
    - 若同目录下存在 requirements.txt 且未强制重生成，则直接安装
    - 否则使用 pipreqs 生成 requirements.txt 后再安装

    参数：
        - **target_path**: 远程的 Python 文件路径或目录路径
        - **index_url**: 可选的 pip 源，例如 https://pypi.tuna.tsinghua.edu.cn/simple
        - **force_regenerate**: 是否强制用 pipreqs 重新生成 requirements.txt
    
    返回：
        str: 操作输出（包含 stdout 与 stderr 合并内容）
    """
    logger.info(f"安装Python依赖: {target_path}")
    # 构建远程执行脚本，合并 stderr 到 stdout 避免将告警误判为错误
    index_part = f" -i {index_url}" if index_url else ""
    force_flag = "true" if force_regenerate else "false"
    # 获取conda环境初始化代码
    conda_init = _get_conda_init_code()
    conda_deinit = _get_conda_deactivate_code()
    # 使用 /bin/sh，确保目录解析与条件判断可用
    script_parts = [
        "set -e; ",
        # 激活conda环境（如果配置了）
        conda_init,
        # 解析目录：如果是目录则直接使用，否则取父目录
        "TP=\"" + target_path.replace("\"", "\\\"") + "\"; ",
        "if [ -d \"$TP\" ]; then DIR=\"$TP\"; else DIR=\"$(dirname \"$TP\")\"; fi; ",
        "REQ=\"$DIR/requirements.txt\"; ",
        # 检测Python命令：在conda环境中使用python，否则使用python3
        "if [ -n \"$CONDA_DEFAULT_ENV\" ] || [ -n \"$(conda info --envs 2>/dev/null | grep -E '^\\*')\" ]; then ",
        "  PYTHON_CMD=\"python\"; ",
        "else ",
        "  PYTHON_CMD=\"python3\"; ",
        "fi; ",
        "PIP=\"$PYTHON_CMD -m pip\"; ",
        # 非 root 时使用 --user，root 下不加 --user
        "if [ \"$(id -u)\" -ne 0 ]; then USER_FLAG=\" --user\"; else USER_FLAG=\"\"; fi; ",
        f"INDEX=\"{index_part}\"; ",
        f"FORCE=\"{force_flag}\"; ",
        # 不升级系统 pip，避免 rpm 管理的 pip 卸载失败
        # 如果已有 requirements 且不强制生成，则直接安装
        "if [ -f \"$REQ\" ] && [ \"$FORCE\" != \"true\" ]; then ",
        "  echo \"Using existing requirements.txt\"; ",
        "  $PIP install -r \"$REQ\"$INDEX$USER_FLAG --no-cache-dir 2>&1; ",
        "else ",
        "  echo \"Generating requirements.txt via pipreqs\"; ",
        "  $PIP install pipreqs$INDEX$USER_FLAG --no-cache-dir 2>&1; ",
        "  cd \"$DIR\"; ",
        "  pipreqs . --force --encoding=utf-8 2>&1; ",
        "  $PIP install -r requirements.txt$INDEX$USER_FLAG --no-cache-dir 2>&1; ",
        "fi; ",
        "echo \"Done.\"; ",
        # 退出conda环境（如果配置了）
        conda_deinit,
    ]
    script = "".join(script_parts)
    # 阻塞的 SSH 调用放到线程池中执行
    result = await asyncio.to_thread(_execute_remote_command, f"sh -lc '{script}'")
    logger.info(f"安装Python依赖结果: {result}")
    return result


@mcp.tool()
async def install_python_deps_async(target_path: str, index_url: Optional[str] = None, force_regenerate: bool = False) -> str:
    """
    异步安装依赖：在远程后台执行依赖安装流程并立即返回 PID 与日志路径。
    之后可使用 tail_log 查看进度，避免 MCP 客户端请求超时。

    参数：
        - **target_path**: 远程的 Python 文件路径或目录路径
        - **index_url**: 可选的 pip 源
        - **force_regenerate**: 是否强制重新生成 requirements.txt

    返回：
        str: "PID:12345 LOG:/tmp/mcp_deps_1699999999_12345.log"
    """
    logger.info(f"异步安装Python依赖: {target_path}")
    # 构建与 install_python_deps 相同的核心脚本，但通过 nohup 后台运行并将输出重定向到日志
    index_part = f" -i {index_url}" if index_url else ""
    force_flag = "true" if force_regenerate else "false"
    # 获取conda环境初始化代码
    conda_init = _get_conda_init_code()
    conda_deinit = _get_conda_deactivate_code()
    script_core_parts = [
        "set -e; ",
        # 激活conda环境（如果配置了）
        conda_init,
        "TP=\"" + target_path.replace("\"", "\\\"") + "\"; ",
        "if [ -d \"$TP\" ]; then DIR=\"$TP\"; else DIR=\"$(dirname \"$TP\")\"; fi; ",
        "REQ=\"$DIR/requirements.txt\"; ",
        # 检测Python命令：在conda环境中使用python，否则使用python3
        "if [ -n \"$CONDA_DEFAULT_ENV\" ] || [ -n \"$(conda info --envs 2>/dev/null | grep -E '^\\*')\" ]; then ",
        "  PYTHON_CMD=\"python\"; ",
        "else ",
        "  PYTHON_CMD=\"python3\"; ",
        "fi; ",
        "PIP=\"$PYTHON_CMD -m pip\"; ",
        "if [ \"$(id -u)\" -ne 0 ]; then USER_FLAG=\" --user\"; else USER_FLAG=\"\"; fi; ",
        f"INDEX=\"{index_part}\"; ",
        f"FORCE=\"{force_flag}\"; ",
        "if [ -f \"$REQ\" ] && [ \"$FORCE\" != \"true\" ]; then ",
        "  echo \"Using existing requirements.txt\"; ",
        "  $PIP install -r \"$REQ\"$INDEX$USER_FLAG --no-cache-dir 2>&1; ",
        "else ",
        "  echo \"Generating requirements.txt via pipreqs\"; ",
        "  $PIP install pipreqs$INDEX$USER_FLAG --no-cache-dir 2>&1; ",
        "  cd \"$DIR\"; ",
        "  pipreqs . --force --encoding=utf-8 2>&1; ",
        "  $PIP install -r requirements.txt$INDEX$USER_FLAG --no-cache-dir 2>&1; ",
        "fi; ",
        "echo \"Done.\"; ",
        # 退出conda环境（如果配置了）
        conda_deinit,
    ]
    script_core = "".join(script_core_parts)
    wrapper = (
        "LOG=\\\"/tmp/mcp_deps_$(date +%s)_$$.log\\\"; "
        f"nohup sh -lc '{script_core}' > \"$LOG\" 2>&1 & echo PID:$! LOG:$LOG"
    )
    # 远程命令同样放到线程池，避免阻塞事件循环
    result = await asyncio.to_thread(_execute_remote_command, f"sh -lc '{wrapper}'")
    logger.info(f"异步安装Python依赖结果: {result}")
    return result


@mcp.tool()
async def tail_log(file_path: str, lines: int = 100) -> str:
    """
    查看远程日志文件的末尾若干行（默认100行）。
    
    参数：
        - **file_path**: 日志文件路径
        - **lines**: 要显示的行数，默认100行
    
    返回：
        str: 日志内容
    """
    logger.info(f"查看日志: {file_path}, 行数: {lines}")
    # 简单转义单引号
    safe = file_path.replace("'", "'\\''")
    # tail 命令通过 SSH 执行是阻塞的，放到线程池
    cmd = f"sh -lc 'tail -n {lines} \"{safe}\" || echo \"(no log or no permission)\"'"
    result = await asyncio.to_thread(_execute_remote_command, cmd)
    logger.info(f"查看日志结果: {result}")
    return result


@mcp.tool()
async def run_python_file(file_path: str, python_bin: Optional[str] = None, args: str = "") -> str:
    """
    在远程服务器执行指定的 Python 文件，并返回标准输出与标准错误（合并）。

    参数：
        - **file_path**: 远程 Python 文件路径
        - **python_bin**: 远程使用的 Python 解释器（None时自动检测：conda环境中使用python，否则使用python3）
        - **args**: 传给脚本的参数字符串（可为空）

    返回：
        str: 合并后的输出文本，包含退出码（ExitCode: N）
    """
    logger.info(f"执行Python文件: {file_path}, 解释器: {python_bin}, 参数: {args}")
    # 使用 /bin/sh -lc 保持 PATH 与环境变量加载，与交互式一致
    # 将 stderr 合并到 stdout，最后打印退出码便于上层判断
    safe_file = file_path.replace("'", "'\\''")
    safe_args = args if args else ""
    # 获取conda环境初始化代码
    conda_init = _get_conda_init_code()
    conda_deinit = _get_conda_deactivate_code()
    
    # 如果未指定python_bin，则在脚本中动态检测
    if python_bin is None:
        # 构建命令，包括conda环境激活和退出，动态检测python命令
        cmd_parts = [
            "sh -lc '",
            conda_init,  # 激活conda环境（如果配置了）
            f"FILE=\"{safe_file}\"; ",
            "DIR=\"$(dirname \"$FILE\")\"; ",
            "BASE=\"$(basename \"$FILE\")\"; ",
            "cd \"$DIR\" || exit 1; ",
            # 检测Python命令：在conda环境中使用python，否则使用python3
            "if [ -n \"$CONDA_DEFAULT_ENV\" ] || [ -n \"$(conda info --envs 2>/dev/null | grep -E '^\\*')\" ]; then ",
            "  PYTHON_CMD=\"python\"; ",
            "else ",
            "  PYTHON_CMD=\"python3\"; ",
            "fi; ",
            "$PYTHON_CMD \"./$BASE\" " + safe_args + " 2>&1; echo ExitCode:$?; ",
            conda_deinit,  # 退出conda环境（如果配置了）
            "'"
        ]
    else:
        # 使用指定的python命令
        cmd_parts = [
            "sh -lc '",
            conda_init,  # 激活conda环境（如果配置了）
            f"FILE=\"{safe_file}\"; ",
            "DIR=\"$(dirname \"$FILE\")\"; ",
            "BASE=\"$(basename \"$FILE\")\"; ",
            "cd \"$DIR\" || exit 1; ",
            f"{python_bin} \"./$BASE\" {safe_args} 2>&1; echo ExitCode:$?; ",
            conda_deinit,  # 退出conda环境（如果配置了）
            "'"
        ]
    cmd = "".join(cmd_parts)
    # 通过 SSH 远程执行 Python 文件是阻塞操作，放到线程池
    result = await asyncio.to_thread(_execute_remote_command, cmd)
    logger.info(f"执行Python文件结果: {result}")
    return result


@mcp.tool()
async def upload_file_to_dir(local_path: str, remote_dir: str, filename: Optional[str] = None, permissions: Optional[str] = None, content_b64: Optional[str] = None) -> str:
    """
    上传文件到远程 Linux 指定目录。

    参数：
        - **local_path**: 本地文件绝对/相对路径；当提供 content_b64 时可为空字符串
        - **remote_dir**: 远程目标目录（不存在将自动创建）
        - **filename**: 远程保存文件名（可选，不填则沿用本地文件名）
        - **permissions**: 可选的八进制权限字符串，例如 "644" 或 "755"
        - **content_b64**: 可选，Base64 编码后的文件二进制内容。提供时将走流式写入路径

    返回：
        str: 成功时返回远程文件完整路径；失败时返回以 "ERROR:" 开头的错误信息
    """
    logger.info(f"上传文件: {local_path} -> {remote_dir}")
    mode = None
    if permissions:
        try:
            mode = int(permissions, 8)
        except ValueError:
            return f"ERROR: 非法权限值: {permissions}"
    # 优先使用流式内容上传
    if content_b64:
        try:
            content = base64.b64decode(content_b64)
        except Exception as e:
            return f"ERROR: Base64 解码失败: {str(e)}"
        # 流式上传必须指定远程文件名
        if not filename:
            return "ERROR: 使用 content_b64 上传时必须提供 filename"
        # SFTP 上传是阻塞操作，放入线程池
        result = await asyncio.to_thread(_sftp_upload_bytes, content, remote_dir, filename, mode)
    else:
        # 否则走本地文件上传，同样放入线程池
        result = await asyncio.to_thread(_sftp_upload_file, local_path, remote_dir, filename, mode)
    
    logger.info(f"上传文件结果: {result}")
    return result


@mcp.tool()
async def delete_file(file_path: str, force: bool = False) -> str:
    """
    删除远程 Linux 服务器上的文件。

    参数：
        - **file_path**: 要删除的文件路径（绝对或相对路径）
        - **force**: 是否强制删除（忽略不存在的文件错误）

    返回：
        str: 操作结果字符串
    """
    logger.info(f"删除文件: {file_path}, 强制模式: {force}")
    
    # 构建删除命令
    if force:
        # 使用 rm -f 强制删除，忽略不存在的文件
        command = f"rm -f '{file_path}'"
    else:
        # 使用 rm 正常删除，如果文件不存在会报错
        command = f"rm '{file_path}'"
    
    # 远程执行删除命令是阻塞的，放到线程池
    result = await asyncio.to_thread(_execute_remote_command, command)
    
    # 检查结果
    if result.startswith("ERROR:"):
        logger.error(f"删除文件失败: {result}")
        return result
    else:
        logger.info(f"删除文件成功: {file_path}")
        return f"成功删除文件: {file_path}"


@mcp.tool()
async def delete_directory(dir_path: str, recursive: bool = False, force: bool = False) -> str:
    """
    删除远程 Linux 服务器上的目录。

    参数：
        - **dir_path**: 要删除的目录路径（绝对或相对路径）
        - **recursive**: 是否递归删除目录及其所有内容
        - **force**: 是否强制删除（忽略不存在的目录错误）

    返回：
        str: 操作结果字符串
    """
    logger.info(f"删除目录: {dir_path}, 递归模式: {recursive}, 强制模式: {force}")
    
    # 构建删除命令
    if recursive:
        if force:
            # 使用 rm -rf 递归强制删除
            command = f"rm -rf '{dir_path}'"
        else:
            # 使用 rm -r 递归删除
            command = f"rm -r '{dir_path}'"
    else:
        if force:
            # 使用 rmdir 删除空目录，但先检查是否存在
            command = f"rmdir '{dir_path}' 2>/dev/null || echo '目录不存在或不为空'"
        else:
            # 使用 rmdir 删除空目录
            command = f"rmdir '{dir_path}'"
    
    # 同样把阻塞的远程删除命令放到线程池
    result = await asyncio.to_thread(_execute_remote_command, command)
    
    # 检查结果
    if result.startswith("ERROR:"):
        logger.error(f"删除目录失败: {result}")
        return result
    else:
        logger.info(f"删除目录成功: {dir_path}")
        return f"成功删除目录: {dir_path}"


def _convert_markdown_to_pdf_with_pypandoc(markdown_content: str, output_path: str) -> str:
    """
    使用pypandoc将markdown转换为PDF
    
    参数：
        - markdown_content: markdown内容
        - output_path: 输出PDF文件路径
    
    返回：
        str: 操作结果
    """
    try:
        # 使用pypandoc转换，支持图表和复杂格式
        # 尝试使用xelatex（支持中文），如果失败则尝试pdflatex
        extra_args = [
            '--variable=geometry:margin=2.5cm',  # 页边距
            '--highlight-style=tango',  # 代码高亮样式
        ]
        
        # 尝试使用xelatex（更好的中文支持）
        try:
            pypandoc.convert_text(
                markdown_content,
                'pdf',
                format='markdown',
                outputfile=output_path,
                extra_args=extra_args + [
                    '--pdf-engine=xelatex',
                    '--variable=mainfont:"Microsoft YaHei"',
                ]
            )
        except Exception:
            # 如果xelatex失败，尝试使用pdflatex
            logger.warning("xelatex转换失败，尝试使用pdflatex")
            pypandoc.convert_text(
                markdown_content,
                'pdf',
                format='markdown',
                outputfile=output_path,
                extra_args=extra_args + [
                    '--pdf-engine=pdflatex',
                ]
            )
        
        return f"成功: PDF已生成到 {output_path}"
    except Exception as e:
        return f"ERROR: pypandoc转换失败: {str(e)}。提示：可能需要安装pandoc和LaTeX引擎（如xelatex或pdflatex）"


def _convert_markdown_to_pdf_with_weasyprint(markdown_content: str, output_path: str) -> str:
    """
    使用weasyprint将markdown转换为PDF（备选方案）
    
    参数：
        - markdown_content: markdown内容
        - output_path: 输出PDF文件路径
    
    返回：
        str: 操作结果
    """
    try:
        # 动态导入weasyprint（如果启动时导入失败）
        try:
            from weasyprint import HTML
        except ImportError:
            # 如果导入失败，尝试重新导入
            try:
                import weasyprint
                HTML = weasyprint.HTML
            except Exception as e:
                return f"ERROR: weasyprint库无法导入: {str(e)}。在Windows上需要安装GTK+库。"
        
        if not MARKDOWN_AVAILABLE:
            return "ERROR: markdown库未安装"
        
        # 将markdown转换为HTML
        md = markdown.Markdown(extensions=[
            'codehilite',
            'tables',
            'fenced_code',
            'nl2br',
            'sane_lists'
        ])
        html_content = md.convert(markdown_content)
        
        # 添加CSS样式
        html_with_style = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                @page {{
                    size: A4;
                    margin: 2.5cm;
                }}
                body {{
                    font-family: "Microsoft YaHei", "SimSun", Arial, sans-serif;
                    line-height: 1.6;
                    color: #333;
                }}
                h1, h2, h3, h4, h5, h6 {{
                    color: #2c3e50;
                    margin-top: 1.5em;
                    margin-bottom: 0.5em;
                }}
                code {{
                    background-color: #f4f4f4;
                    padding: 2px 4px;
                    border-radius: 3px;
                    font-family: "Consolas", "Monaco", monospace;
                }}
                pre {{
                    background-color: #f8f8f8;
                    border: 1px solid #ddd;
                    border-radius: 5px;
                    padding: 1em;
                    overflow-x: auto;
                }}
                table {{
                    border-collapse: collapse;
                    width: 100%;
                    margin: 1em 0;
                }}
                th, td {{
                    border: 1px solid #ddd;
                    padding: 8px;
                    text-align: left;
                }}
                th {{
                    background-color: #4CAF50;
                    color: white;
                }}
                img {{
                    max-width: 100%;
                    height: auto;
                }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
        
        # 使用weasyprint生成PDF
        HTML(string=html_with_style).write_pdf(output_path)
        return f"成功: PDF已生成到 {output_path}"
    except Exception as e:
        return f"ERROR: weasyprint转换失败: {str(e)}"


def _convert_markdown_to_word_with_pypandoc(markdown_content: str, output_path: str) -> str:
    """
    使用pypandoc将markdown转换为Word文档
    
    参数：
        - markdown_content: markdown内容
        - output_path: 输出Word文件路径
    
    返回：
        str: 操作结果
    """
    try:
        pypandoc.convert_text(
            markdown_content,
            'docx',
            format='markdown',
            outputfile=output_path,
            extra_args=[
                '--highlight-style=tango',
            ]
        )
        return f"成功: Word文档已生成到 {output_path}"
    except Exception as e:
        return f"ERROR: pypandoc转换失败: {str(e)}"


def _convert_markdown_to_word_with_python_docx(markdown_content: str, output_path: str) -> str:
    """
    使用python-docx将markdown转换为Word文档（备选方案）
    
    参数：
        - markdown_content: markdown内容
        - output_path: 输出Word文件路径
    
    返回：
        str: 操作结果
    """
    try:
        if not PYTHON_DOCX_AVAILABLE:
            return "ERROR: python-docx库未安装"
        if not MARKDOWN_AVAILABLE:
            return "ERROR: markdown库未安装"
        
        # 创建Word文档
        doc = Document()
        
        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Microsoft YaHei'
        font.size = Pt(11)
        
        # 将markdown转换为HTML然后解析
        md = markdown.Markdown(extensions=['codehilite', 'tables', 'fenced_code'])
        html_content = md.convert(markdown_content)
        
        # 简单的HTML解析（处理基本元素）
        lines = markdown_content.split('\n')
        in_code_block = False
        code_block_lines = []
        
        for line in lines:
            # 处理代码块
            if line.strip().startswith('```'):
                if in_code_block:
                    # 结束代码块
                    if code_block_lines:
                        p = doc.add_paragraph()
                        run = p.add_run('\n'.join(code_block_lines))
                        run.font.name = 'Consolas'
                        run.font.size = Pt(10)
                    code_block_lines = []
                    in_code_block = False
                else:
                    in_code_block = True
                continue
            
            if in_code_block:
                code_block_lines.append(line)
                continue
            
            # 处理标题
            if line.startswith('# '):
                doc.add_heading(line[2:].strip(), level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith('#### '):
                doc.add_heading(line[5:].strip(), level=4)
            # 处理列表
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                doc.add_paragraph(line.strip()[2:], style='List Bullet')
            elif re.match(r'^\d+\.\s', line.strip()):
                doc.add_paragraph(re.sub(r'^\d+\.\s', '', line.strip()), style='List Number')
            # 处理普通段落
            elif line.strip():
                doc.add_paragraph(line.strip())
        
        # 保存文档
        doc.save(output_path)
        return f"成功: Word文档已生成到 {output_path}"
    except Exception as e:
        return f"ERROR: python-docx转换失败: {str(e)}"


def _read_remote_file(file_path: str) -> str:
    """
    读取远程文件内容（文本模式）
    
    参数：
        - file_path: 远程文件路径
    
    返回：
        str: 文件内容，失败时返回以ERROR:开头的错误信息
    """
    client = paramiko.SSHClient()
    client.set_missing_host_key_policy(paramiko.AutoAddPolicy())
    try:
        client.connect(REMOTE_HOST, username=USERNAME, password=PASSWORD, timeout=30)
        sftp = client.open_sftp()
        with sftp.file(file_path, mode='r') as f:
            content = f.read().decode('utf-8', errors='ignore')
        sftp.close()
        return content
    except Exception as e:
        return f"ERROR: 读取远程文件失败: {str(e)}"
    finally:
        client.close()


def _read_remote_file_bytes(file_path: str) -> bytes:
    """
    读取远程文件内容（二进制模式）
    
    参数：
        - file_path: 远程文件路径
    
    返回：
        bytes: 文件内容，失败时返回None
    """
    client = paramiko.SSHClient()
    client.set_missing_host_key_policy(paramiko.AutoAddPolicy())
    try:
        client.connect(REMOTE_HOST, username=USERNAME, password=PASSWORD, timeout=30)
        sftp = client.open_sftp()
        with sftp.file(file_path, mode='rb') as f:
            content = f.read()
        sftp.close()
        return content
    except Exception as e:
        logger.error(f"读取远程文件失败: {str(e)}")
        return None
    finally:
        client.close()


def _ensure_remote_directory(dir_path: str) -> bool:
    """
    确保远程目录存在
    
    参数：
        - dir_path: 远程目录路径
    
    返回：
        bool: 是否成功
    """
    try:
        result = _execute_remote_command(f"mkdir -p '{dir_path}'")
        return not result.startswith("ERROR:")
    except Exception as e:
        logger.error(f"创建远程目录失败: {str(e)}")
        return False


@mcp.tool()
async def convert_markdown_to_document(
    markdown_content: str,
    output_format: str = "pdf"
) -> str:
    """
    将标准的markdown文档转换为优雅的PDF或Word文档，支持图表、代码块、表格等内容。
    
    参数：
        - **markdown_content**: markdown内容字符串（必需）
        - **output_format**: 输出格式，可选 "pdf" 或 "docx"（默认: "pdf"）
    
    返回：
        str: 操作结果，成功时返回文件路径和下载URL，失败时返回以ERROR:开头的错误信息
    
    支持的markdown特性：
        - 标题、段落、列表
        - 代码块（带语法高亮）
        - 表格
        - 图片（需要图片文件在远程服务器可访问）
        - 链接
        - 粗体、斜体等格式
    """
    logger.info(f"转换markdown文档，格式: {output_format}, 内容长度: {len(markdown_content)}")
    
    try:
        # 验证markdown内容
        if not markdown_content or not markdown_content.strip():
            return "ERROR: markdown_content 不能为空"
        
        md_content = markdown_content
        logger.info(f"使用提供的markdown内容，长度: {len(md_content)}")
        
        # 确保输出目录存在（本地目录）
        os.makedirs(DOCUMENT_OUTPUT_DIR, exist_ok=True)
        
        # 自动生成输出文件名（使用时间戳确保唯一性）
        timestamp = int(time.time())
        filename = f"markdown_export_{timestamp}.{output_format}"
        output_path = os.path.join(DOCUMENT_OUTPUT_DIR, filename)
        
        logger.info(f"输出路径: {output_path}")
        
        # 在本地临时目录生成文件
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{output_format}") as tmp_file:
            local_temp_path = tmp_file.name
        
        try:
            # 根据格式选择转换方法
            result = None
            if output_format.lower() == "pdf":
                # 优先使用pypandoc，失败时自动回退到weasyprint
                if PYPANDOC_AVAILABLE:
                    logger.info("尝试使用pypandoc转换PDF")
                    result = await asyncio.to_thread(_convert_markdown_to_pdf_with_pypandoc, md_content, local_temp_path)
                    # 如果pypandoc失败（通常是pandoc未安装），尝试备选方案
                    if result.startswith("ERROR:"):
                        error_lower = result.lower()
                        if any(keyword in error_lower for keyword in ["pandoc", "no pandoc", "install pandoc"]):
                            logger.warning(f"pypandoc转换失败（pandoc未安装），自动切换到weasyprint")
                            result = None  # 清空结果，尝试备选方案
                        else:
                            # 其他错误也尝试备选方案
                            logger.warning(f"pypandoc转换失败，尝试使用备选方案")
                            result = None
                
                # 如果pypandoc不可用或失败，使用weasyprint
                if not result or result.startswith("ERROR:"):
                    # 动态检测weasyprint是否可用（即使导入时失败，运行时可能可用）
                    weasyprint_usable = False
                    if WEASYPRINT_AVAILABLE:
                        weasyprint_usable = True
                    else:
                        # 尝试动态导入weasyprint
                        try:
                            from weasyprint import HTML
                            weasyprint_usable = True
                            logger.info("动态检测到weasyprint可用")
                        except Exception:
                            weasyprint_usable = False
                    
                    if weasyprint_usable:
                        logger.info("使用weasyprint转换PDF（备选方案）")
                        result = await asyncio.to_thread(_convert_markdown_to_pdf_with_weasyprint, md_content, local_temp_path)
                    else:
                        return "ERROR: 未安装可用的PDF转换工具。请安装: pip install pypandoc（需要系统安装pandoc）或 pip install weasyprint（Windows上需要GTK+库）"
            
            elif output_format.lower() in ["docx", "word"]:
                # 优先使用pypandoc，失败时自动回退到python-docx
                if PYPANDOC_AVAILABLE:
                    logger.info("尝试使用pypandoc转换Word")
                    result = await asyncio.to_thread(_convert_markdown_to_word_with_pypandoc, md_content, local_temp_path)
                    # 如果pypandoc失败（通常是pandoc未安装），尝试备选方案
                    if result.startswith("ERROR:"):
                        error_lower = result.lower()
                        if any(keyword in error_lower for keyword in ["pandoc", "no pandoc", "install pandoc"]):
                            logger.warning(f"pypandoc转换失败（pandoc未安装），自动切换到python-docx")
                            result = None  # 清空结果，尝试备选方案
                        else:
                            # 其他错误也尝试备选方案
                            logger.warning(f"pypandoc转换失败，尝试使用备选方案")
                            result = None
                
                # 如果pypandoc不可用或失败，使用python-docx
                if not result or result.startswith("ERROR:"):
                    if PYTHON_DOCX_AVAILABLE:
                        logger.info("使用python-docx转换Word（备选方案）")
                        result = await asyncio.to_thread(_convert_markdown_to_word_with_python_docx, md_content, local_temp_path)
                    else:
                        return "ERROR: 未安装可用的Word转换工具。请安装: pip install pypandoc（需要系统安装pandoc）或 pip install python-docx"
            else:
                return f"ERROR: 不支持的输出格式: {output_format}，支持格式: pdf, docx"
            
            if result.startswith("ERROR:"):
                logger.error(f"转换失败: {result}")
                return result
            
            # 验证本地文件是否已创建
            if not os.path.exists(local_temp_path):
                return f"ERROR: 文件转换完成但未找到本地临时文件: {local_temp_path}"
            
            # 直接复制到目标路径（本地保存）
            shutil.copy2(local_temp_path, output_path)
            
            # 验证文件是否已创建
            if not os.path.exists(output_path):
                return f"ERROR: 文件保存失败: {output_path}"
            
        finally:
            # 清理本地临时文件
            try:
                if os.path.exists(local_temp_path):
                    os.unlink(local_temp_path)
            except Exception as e:
                logger.warning(f"清理临时文件失败: {str(e)}")
        
        # 获取文件大小（本地文件）
        try:
            file_size = os.path.getsize(output_path)
        except Exception as e:
            logger.warning(f"获取文件大小失败: {str(e)}")
            file_size = 0
        
        # 生成下载URL（使用正斜杠）
        filename = os.path.basename(output_path)
        # 将Windows路径转换为URL格式（使用正斜杠）
        url_path = output_path.replace('\\', '/')
        download_url = f"http://{HTTP_SERVER_HOST}:8002/download/document?file_path={url_path}"
        
        # 如果转换成功，返回输出文件路径和下载信息
        logger.info(f"转换成功: {result}")
        return f"转换成功！\n输出文件: {output_path}\n文件大小: {file_size} 字节\n下载URL: {download_url}\n文件名: {filename}"
        
    except Exception as e:
        error_msg = f"ERROR: 转换过程中发生异常: {str(e)}"
        logger.error(error_msg, exc_info=True)
        return error_msg


# HTTP接口定义
@app.post("/upload/risk-analysis")
async def upload_file_to_risk_analysis(
    userId: Optional[str] = Form(None),
    file: UploadFile = File(...)
):
    """
    上传文件到风险分析目录
    
    参数：
        - **file**: 上传的文件
    
    返回：
        JSON响应，包含上传结果
    """
    logger.info(f"收到文件上传请求: {file.filename}")
    
    # 验证文件格式
    is_valid, error_msg = _validate_file_format(file.filename)
    if not is_valid:
        logger.error(f"文件格式验证失败: {error_msg}")
        raise HTTPException(status_code=400, detail=error_msg)
    
    # 读取文件内容
    try:
        file_content = await file.read()
        logger.info(f"文件读取成功，大小: {len(file_content)} 字节")
    except Exception as e:
        logger.error(f"文件读取失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"文件读取失败: {str(e)}")
    
    # 获取原始文件扩展名
    original_ext = os.path.splitext(file.filename)[1]
    sanitized_user_id = userId.strip() if userId else ""
    if sanitized_user_id:
        new_filename = f"data_{sanitized_user_id}{original_ext}"
    else:
        new_filename = f"data{original_ext}"
    
    # 上传到远程服务器
    try:
        # 使用现有的_sftp_upload_bytes函数上传文件
        result = _sftp_upload_bytes(
            content=file_content,
            remote_dir=RISK_ANALYSIS_DIR,
            filename=new_filename,
            mode=0o644
        )
        
        if result.startswith("ERROR:"):
            logger.error(f"文件上传失败: {result}")
            raise HTTPException(status_code=500, detail=result)
        
        logger.info(f"文件上传成功: {result}")
        return JSONResponse(
            status_code=200,
            content={
                "success": True,
                "message": "文件上传成功",
                "original_filename": file.filename,
                "remote_path": result,
                "file_size": len(file_content)
            }
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"文件上传过程中发生错误: {str(e)}")
        raise HTTPException(status_code=500, detail=f"文件上传失败: {str(e)}")


@app.get("/download/document")
async def download_document(file_path: str = Query(..., description="本地文件路径")):
    """
    下载生成的文档文件
    
    参数：
        - **file_path**: 本地文件路径（必须位于DOCUMENT_OUTPUT_DIR目录下）
    
    返回：
        文件流响应
    """
    logger.info(f"收到文档下载请求: {file_path}")
    
    # 安全检查：确保文件路径在允许的目录下
    try:
        # 处理URL中的路径（可能包含正斜杠）
        normalized_path = file_path.replace('/', os.sep)
        abs_file_path = os.path.abspath(normalized_path)
        abs_output_dir = os.path.abspath(DOCUMENT_OUTPUT_DIR)
        
        # 使用commonpath确保路径在允许的目录下
        try:
            common_path = os.path.commonpath([abs_file_path, abs_output_dir])
            # 规范化路径比较（处理大小写和尾随分隔符）
            if os.path.normpath(common_path) != os.path.normpath(abs_output_dir):
                logger.error(f"非法文件路径访问尝试: {file_path}")
                raise HTTPException(status_code=403, detail="文件路径不在允许的目录下")
        except ValueError:
            # 路径不在同一驱动器或无效
            logger.error(f"非法文件路径访问尝试: {file_path}")
            raise HTTPException(status_code=403, detail="文件路径不在允许的目录下")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"路径安全检查失败: {str(e)}")
        raise HTTPException(status_code=403, detail="文件路径不在允许的目录下")
    
    # 读取本地文件
    if not os.path.exists(normalized_path):
        logger.error(f"文件不存在: {normalized_path}")
        raise HTTPException(status_code=404, detail="文件不存在")
    
    try:
        with open(normalized_path, 'rb') as f:
            file_content = f.read()
    except Exception as e:
        logger.error(f"文件读取失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"文件读取失败: {str(e)}")
    
    # 获取文件名和MIME类型
    filename = os.path.basename(file_path)
    file_ext = os.path.splitext(filename)[1].lower()
    
    # 根据文件扩展名设置MIME类型
    mime_types = {
        '.pdf': 'application/pdf',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.doc': 'application/msword',
    }
    media_type = mime_types.get(file_ext, 'application/octet-stream')
    
    logger.info(f"文件下载成功: {filename}, 大小: {len(file_content)} 字节")
    
    # 返回文件流
    return StreamingResponse(
        iter([file_content]),
        media_type=media_type,
        headers={
            "Content-Disposition": f"attachment; filename={filename}",
            "Content-Length": str(len(file_content))
        }
    )


@app.get("/documents/list")
async def list_documents():
    """
    列出所有生成的文档文件
    
    返回：
        JSON响应，包含文件列表
    """
    logger.info("收到文档列表请求")
    
    try:
        # 列出本地目录下的所有文件
        files = []
        if not os.path.exists(DOCUMENT_OUTPUT_DIR):
            logger.warning(f"输出目录不存在: {DOCUMENT_OUTPUT_DIR}")
            return JSONResponse(
                status_code=200,
                content={
                    "success": True,
                    "count": 0,
                    "files": []
                }
            )
        
        # 遍历目录查找PDF和Word文件
        for root, dirs, filenames in os.walk(DOCUMENT_OUTPUT_DIR):
            for filename in filenames:
                file_ext = os.path.splitext(filename)[1].lower()
                if file_ext in ['.pdf', '.docx', '.doc']:
                    file_path = os.path.join(root, filename)
                    try:
                        # 获取文件信息
                        stat_info = os.stat(file_path)
                        file_size = stat_info.st_size
                        mtime = int(stat_info.st_mtime)
                        
                        # 转换为标准路径格式（使用正斜杠用于URL）
                        normalized_path = file_path.replace('\\', '/')
                        
                        files.append({
                            "filename": filename,
                            "path": normalized_path,
                            "size": file_size,
                            "modified_time": mtime,
                            "format": file_ext[1:] if file_ext else "unknown",
                            "download_url": f"http://{HTTP_SERVER_HOST}:8002/download/document?file_path={normalized_path}"
                        })
                    except Exception as e:
                        logger.warning(f"获取文件信息失败 {file_path}: {str(e)}")
                        continue
        
        # 按修改时间倒序排序
        files.sort(key=lambda x: x['modified_time'], reverse=True)
        
        return JSONResponse(
            status_code=200,
            content={
                "success": True,
                "count": len(files),
                "files": files
            }
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"列出文档失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"列出文档失败: {str(e)}")


@app.get("/health")
async def health_check():
    """健康检查接口"""
    return JSONResponse(
        status_code=200,
        content={
            "status": "healthy",
            "service": "Linux File Operations API",
            "version": "1.0.0"
        }
    )


def run_http_server():
    """在单独线程中运行HTTP服务器"""
    try:
        logger.info("启动HTTP服务器...")
        uvicorn.run(app, host="0.0.0.0", port=8002, log_level="info")
    except Exception as e:
        logger.error(f"HTTP服务器启动失败: {e}")


def log_thread_statuses():
    """输出当前进程中所有线程的状态，便于排查卡死问题"""
    threads = threading.enumerate()
    logger.info(f"[线程监控] 当前线程数: {len(threads)}")
    for t in threads:
        logger.info(
            "[线程监控] 名称=%s, ident=%s, daemon=%s, alive=%s",
            t.name,
            t.ident,
            t.daemon,
            t.is_alive()
        )


def thread_status_monitor():
    """后台循环监控线程状态"""
    while True:
        try:
            log_thread_statuses()
        except Exception as exc:
            logger.warning(f"[线程监控] 输出线程状态失败: {exc}")
        time.sleep(THREAD_STATUS_INTERVAL)


if __name__ == "__main__":
    try:
        logger.info("正在启动 Linux File Operations 服务器...")
        logger.info("MCP服务器将监听: 0.0.0.0:8001")
        logger.info("HTTP服务器将监听: 0.0.0.0:8002")
        logger.info("使用 SSE 传输协议")
        logger.info("SSE 端点: http://localhost:8001/sse")
        logger.info("使用本地docker启动的dify使用：SSE端点: http://host.docker.internal:8001/sse")
        logger.info(f"HTTP API 端点: http://{HTTP_SERVER_HOST}:8002")
        logger.info(f"文件上传接口: POST http://{HTTP_SERVER_HOST}:8002/upload/risk-analysis")
        logger.info(f"健康检查接口: GET http://{HTTP_SERVER_HOST}:8002/health")
        logger.info("可用MCP工具:")
        logger.info("- create_file: 创建文件")
        logger.info("- write_file: 写入文件")
        logger.info("- chmod: 修改文件权限")
        logger.info("- install_python_deps: 安装Python依赖")
        logger.info("- install_python_deps_async: 异步安装Python依赖")
        logger.info("- tail_log: 查看日志")
        logger.info("- run_python_file: 执行Python文件")
        logger.info("- upload_file_to_dir: 上传文件")
        logger.info("- delete_file: 删除文件")
        logger.info("- delete_directory: 删除目录")
        logger.info("- convert_markdown_to_document: 将markdown转换为PDF或Word文档")
        
        # 输出markdown转换工具依赖状态
        logger.info("Markdown转换工具依赖状态:")
        logger.info(f"  - pypandoc: {'已安装' if PYPANDOC_AVAILABLE else '未安装（推荐安装以获得最佳效果）'}")
        logger.info(f"  - weasyprint: {'已安装' if WEASYPRINT_AVAILABLE else '未安装（PDF备选方案）'}")
        logger.info(f"  - python-docx: {'已安装' if PYTHON_DOCX_AVAILABLE else '未安装（Word备选方案）'}")
        logger.info(f"  - markdown: {'已安装' if MARKDOWN_AVAILABLE else '未安装（必需）'}")
        logger.info("可用HTTP接口:")
        logger.info("- POST /upload/risk-analysis: 上传文件到风险分析目录")
        logger.info("- GET /download/document: 下载生成的文档文件")
        logger.info("- GET /documents/list: 列出所有生成的文档")
        logger.info("- GET /health: 健康检查")
        
        # 在单独线程中启动HTTP服务器
        http_thread = threading.Thread(target=run_http_server, daemon=True)
        http_thread.start()

        # 启动线程状态监控线程
        monitor_thread = threading.Thread(
            target=thread_status_monitor,
            name="ThreadStatusMonitor",
            daemon=True
        )
        monitor_thread.start()
        
        # 等待HTTP服务器启动
        time.sleep(2)
        
        # 启动MCP服务器（SSE模式）
        mcp.run(transport="sse")
        
    except KeyboardInterrupt:
        logger.info("服务器被用户中断")
    except Exception as e:
        logger.error(f"服务器运行出错: {e}")
        logger.error("请检查端口是否被占用或网络配置")
        raise